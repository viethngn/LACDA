import xlrd
import datetime
import texttable as tt 
from docx import Document

class LACDA_Core:
	date1 = input("Report's start date: ")
	print(date1)
	date2 = input("Report's end date: ")
	print(date2)
	month = input("Report's month: ")
	print(month)
	string = input("Name of the insights xlsx file: ")
	print(string)
	print()

	file = xlrd.open_workbook(string)
	#"Facebook Insights Data Export (Post Level) - LOK - 2017-07-27.xlsx"
	main_sheet = file.sheet_by_name("Key Metrics")
	minor_sheet = file.sheet_by_name("Lifetime Talking About This(...")
	data_col = {}

	data_entry = ["Post ID", 
					"Permalink", 
					"Type", 
					"Posted", 
					"Lifetime: The total number of people your Page post was served to. (Unique Users)", 
					"Lifetime: The number of people your advertised Page post was served to. (Unique Users)", 
					"Lifetime: The number of unique people who engaged in certain ways with your Page post, for example by commenting on, liking, sharing, or clicking upon particular elements of the post. (Unique Users)", 
					"Lifetime: Number of times your video was viewed to 95% of its length without any paid promotion. (Unique Users)", 
					"Lifetime: Number of times your video was viewed to 95% of its length without any paid promotion. (Total Count)", 
					"Lifetime: Number of times your video was viewed to 95% of its length after paid promotion. (Unique Users)", 
					"Lifetime: Number of times your video was viewed to 95% of its length after paid promotion. (Total Count)", 
					"Lifetime: Number of times your video was viewed for more than 3 seconds without any paid promotion. (Unique Users)", 
					"Lifetime: Number of times your video was viewed for more than 3 seconds without any paid promotion. (Total Count)", 
					"Lifetime: Number of times your video was viewed more than 3 seconds after paid promotion. (Unique Users)", 
					"Lifetime: Number of times your video was viewed more than 3 seconds after paid promotion. (Total Count)", 
					"Lifetime: Number of times your video was viewed more than 3 seconds after paid promotion. (Total Count)", 
					"Lifetime: Length of a video post (Total Count)", 
					"Content", 
					"share", 
					"like", 
					"comment"]

	#get column index
	def get_col_index(self, colname, sheetname):
		if sheetname == self.minor_sheet:
			for colindex in range(sheetname.ncols):
				if sheetname.cell(0, colindex).value == colname:
					return colindex	
		else:				
			for colindex in range(sheetname.ncols):
				if colindex > 7:
					if sheetname.cell(1, colindex).value == colname:
						return colindex
				else:
					if sheetname.cell(0, colindex).value == colname:
						return colindex	


	'''for i in range(len(data_entry)):
	 	if data_entry[i] == "share" or data_entry[i] == "like" or data_entry[i] == "comment":
	 		data_col[data_entry[i]] = get_col_index(data_entry[i], self.minor_sheet)
	 	else:
	 		data_col[data_entry[i]] = get_col_index(data_entry[i], self.main_sheet)'''

	#post's likes
	def get_like(self, Permalink):
		like_index= self.get_col_index("like", self.minor_sheet)
		link_index = self.get_col_index("Permalink", self.minor_sheet)
		for index in range(self.minor_sheet.nrows):
			if self.minor_sheet.cell(index, link_index).value == Permalink:
				if self.minor_sheet.cell(index, like_index).value == "":
					return 0
				else:
					return self.minor_sheet.cell(index, like_index).value

	#post's shares
	def get_share(self, Permalink):
		share_index = self.get_col_index("share", self.minor_sheet)
		link_index = self.get_col_index("Permalink", self.minor_sheet)
		for index in range(self.minor_sheet.nrows):
			if self.minor_sheet.cell(index, link_index).value == Permalink:
				if self.minor_sheet.cell(index, share_index).value == "":
					return 0
				else:
					return self.minor_sheet.cell(index, share_index).value

	#post's comments
	def get_comment(self, Permalink):
		cmt_index = self.get_col_index("comment", self.minor_sheet)
		link_index = self.get_col_index("Permalink", self.minor_sheet)
		for index in range(self.minor_sheet.nrows):
			if self.minor_sheet.cell(index, link_index).value == Permalink:
				if self.minor_sheet.cell(index, cmt_index).value == "":
					return 0
				else:
					return self.minor_sheet.cell(index, cmt_index).value

	#post's view
	def get_reach(self, Permalink):
		link_index = self.get_col_index("Permalink", self.main_sheet)
		unpaid_index = self.get_col_index("Lifetime: The total number of people your Page post was served to. (Unique Users)", self.main_sheet)
		paid_index = self.get_col_index("Lifetime: The number of people your advertised Page post was served to. (Unique Users)", self.main_sheet)
		for index in range(self.main_sheet.nrows):
			if self.main_sheet.cell(index, link_index).value == Permalink:
				#if self.main_sheet.cell(index, paid_index).value != 0:
					#return self.main_sheet.cell(index, paid_index).value
				#else:
					return self.main_sheet.cell(index, unpaid_index).value

	#post's reaction ratio
	def get_reaction_ratio(self, Permalink):
		link_index = self.get_col_index("Permalink", self.main_sheet)
		reach = self.get_reach(Permalink)
		ratio = {}
		for index in range(self.main_sheet.nrows):
			if self.main_sheet.cell(index, link_index).value == Permalink:
				ratio['like'] = round((self.get_like(Permalink))*100/reach, 3)
				ratio['share'] = round((self.get_share(Permalink))*100/reach, 3)
				ratio['comment'] = round((self.get_comment(Permalink))*100/reach, 3)
				return ratio

	#topic's reaction and their ratio with reach
	def get_topic_reaction(self, topic):
		arranged = self.arrange_topic()
		if not arranged.has_key(topic):
			return null
		reaction = {'like', 'share', 'comment', 'like ratio', 'shrae ratio', 'comment ratio'}
		total_reach = 0
		total_like = 0
		total_share = 0
		total_comment = 0
		for index in range(len(arranged[topic])):
			total_view += get_reach(arranged[topic][index])
			total_like += get_like(arranged[topic][index])
			total_share += get_share(arranged[topic][index])
			total_comment += get_comment(arranged[topic][index])
		reaction['like'] = total_like
		reaction['share'] = total_share
		reaction['comment'] = total_comment
		reaction['like ratio'] = round(total_like*100/total_reach, 3)
		reaction['share ratio'] = round(total_share*100/total_reach, 3)
		reaction['comment ratio'] = round(total_comment*100/total_reach, 3)
		return reaction

	#list of topic with posts
	def arrange_topic(self):
		topic_index = self.get_col_index("Topic", self.main_sheet)
		link_index = self.get_col_index("Permalink", self.main_sheet)
		arranged = {}
		for index in range(self.main_sheet.nrows):
			if not arranged_topic.has_key(self.main_sheet.cell(index, topic_index)):
				arranged[self.main_sheet.cell(index, topic_index)] = []
				arranged[self.main_sheet.cell(index, topic_index)].append(self.main_sheet.cell(index, link_index))
			else:
				arranged[self.main_sheet.cell(index, topic_index)].append(self.main_sheet.cell(index, link_index))
		return arranged

	#95% views/total views
	def get_audience_ratio(self, Permalink):
		link_index = self.get_col_index("Permalink", self.main_sheet)
		complete_paid_index = self.get_col_index("Lifetime: Number of times your video was viewed to 95% of its length after paid promotion. (Unique Users)", self.main_sheet)
		total_paid_index = self.get_col_index("Lifetime: Number of times your video was viewed more than 3 seconds after paid promotion. (Unique Users)", self.main_sheet)
		complete_unpaid_index = self.get_col_index("Lifetime: Number of times your video was viewed to 95% of its length without any paid promotion. (Unique Users)", self.main_sheet)
		total_unpaid_index = self.get_col_index("Lifetime: Number of times your video was viewed for more than 3 seconds without any paid promotion. (Unique Users)", self.main_sheet)
		for index in range(self.main_sheet.nrows):
			if self.main_sheet.cell(index, link_index).value == Permalink: 				#search for the needed link in the column index Permalink
				if (self.main_sheet.cell(index, complete_paid_index).value != 0) and (self.main_sheet.cell(index, total_paid_index).value != 0):
					audience_watch_full = self.main_sheet.cell(index, complete_paid_index).value
					audience_total = self.main_sheet.cell(index, total_paid_index).value
				elif (self.main_sheet.cell(index, complete_unpaid_index).value != 0) and (self.main_sheet.cell(index, total_unpaid_index ).value != 0):
					audience_watch_full = self.main_sheet.cell(index, complete_unpaid_index).value
					audience_total = self.main_sheet.cell(index, total_unpaid_index).value	
				else:
					return 0	 		
		return round((audience_watch_full)*100 / (audience_total), 3) 

	#average watch duration
	def get_average_duration(self, Permalink):
		link_index = self.get_col_index("Permalink", self.main_sheet)
		average_index = self.get_col_index("Lifetime: Number of times your video was viewed more than 3 seconds after paid promotion. (Total Count)", self.main_sheet)
		for index in range(self.main_sheet.nrows):
			if self.main_sheet.cell(index, link_index).value == Permalink:
				return (float(self.main_sheet.cell(index, average_index).value) / 1000)

	#average duration/total duration
	def get_duration_ratio(self, Permalink):
		link_index = self.get_col_index("Permalink", self.main_sheet)
		average_index = self.get_col_index("Lifetime: Number of times your video was viewed more than 3 seconds after paid promotion. (Total Count)", self.main_sheet)
		length_index = self.get_col_index("Lifetime: Length of a video post (Total Count)", self.main_sheet)
		for index in range(self.main_sheet.nrows):
			if self.main_sheet.cell(index, link_index).value == Permalink:
				if self.main_sheet.cell(index, length_index).value == 0:
					return 0
				else:
					return round((self.main_sheet.cell(index, average_index).value*100 / self.main_sheet.cell(index, length_index).value), 3)

	#promotion efficiency
	def get_promotion_ratio(self, Permalink):
		unique_paid_index = self.get_col_index("Lifetime: Number of times your video was viewed more than 3 seconds after paid promotion. (Unique Users)", self.main_sheet)
		total_paid_index = self.get_col_index("Lifetime: Number of times your video was viewed more than 3 seconds after paid promotion. (Total Count)", self.main_sheet)
		unique_unpaid_index = self.get_col_index("Lifetime: Number of times your video was viewed for more than 3 seconds without any paid promotion. (Unique Users)", self.main_sheet)
		total_unpaid_index = self.get_col_index("Lifetime: Number of times your video was viewed for more than 3 seconds without any paid promotion. (Total Count)", self.main_sheet)
		link_index = self.get_col_index("Permalink", self.main_sheet)
		unpaid_index = self.get_col_index("Lifetime: The total number of people your Page post was served to. (Unique Users)", self.main_sheet)
		paid_index = self.get_col_index("Lifetime: The number of people your advertised Page post was served to. (Unique Users)", self.main_sheet)
		promotion_ratio = {'audience': [], 'watch_time': [], 'reach': [] }
		for index in range(self.main_sheet.nrows):
			if self.main_sheet.cell(index, 1).value == Permalink:
				if self.main_sheet.cell(index, unique_unpaid_index).value == 0:
					promotion_ratio['audience'].append(0)
					promotion_ratio['watch_time'].append(0)
					promotion_ratio['reach'].append(0)
				else:
					promotion_ratio['audience'].append(round(self.main_sheet.cell(index, unique_paid_index).value*100 / (self.main_sheet.cell(index, unique_paid_index).value + self.main_sheet.cell(index, unique_unpaid_index).value), 3))
					promotion_ratio['watch_time'].append(round(self.main_sheet.cell(index, total_paid_index).value*100 / (self.main_sheet.cell(index, total_paid_index).value + self.main_sheet.cell(index, total_unpaid_index).value), 3))
					promotion_ratio['reach'].append(round(self.main_sheet.cell(index, paid_index).value*100 / (self.main_sheet.cell(index, unpaid_index).value + self.main_sheet.cell(index, paid_index).value), 3))
		return promotion_ratio

	def print(self):
		post_index = self.get_col_index("Post Message", self.main_sheet)
		link_index = self.get_col_index("Permalink", self.main_sheet)
		time_index = self.get_col_index("Posted", self.main_sheet)
		type_index = self.get_col_index("Type", self.main_sheet)
		print("LOK POST PERFORMANCE WEEKLY REPORT")
		print("từ " + self.date1 + "/" + self.month + " đến " + self.date2 + "/" + self.month)
		print()
		print("I. Theo thời lượng của post")
		tab1 = tt.Texttable()
		headings1 = ['Post name', 'Hình thức', 'Thời lượng trung bình', 'Trung bình/Full', 'Tỉ lệ xem hết 95%']
		tab1.header(headings1)
		lines1 = []
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			n1 = self.get_average_duration(link)
			n2 = self.get_duration_ratio(link)
			n3 = self.get_audience_ratio(link)
			l1 = []
			l1.append(str(name[:10]))
			l1.append(self.main_sheet.cell(index, type_index).value)
			l1.append(n1)
			l1.append(n2)
			l1.append(n3)
			lines1.append(l1)
		for row in range(len(lines1)):
			tab1.add_row(lines1[row])
		s1 = tab1.draw()
		print(s1)
		print()
		print("II. Theo thời gian post")
		tab2 = tt.Texttable()
		headings2 = ['Post name', 'Hình thức', 'Thời gian post', 'Tỉ lệ like', 'Tỉ lệ share', 'Tỉ lệ comment'] 
		tab2.header(headings2)
		lines2 = []
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			time = self.main_sheet.cell(index, time_index).value[10:]
			ratio = self.get_reaction_ratio(link)
			l2 = []
			l2.append(str(name[:10]))
			l2.append(self.main_sheet.cell(index, type_index).value)
			l2.append(time)
			l2.append(ratio['like'])
			l2.append(ratio['share'])
			l2.append(ratio['comment'])
			lines2.append(l2)
		for row in range(len(lines2)):
			tab2.add_row(lines2[row])
		s2 = tab2.draw()
		print(s2)
		print()
		print("III. Hiệu quả của promotion")
		tab3 = tt.Texttable()
		headings3 = ['Post name', 'Hình thức', 'Lượng khán giả', 'Lượt xem', 'Lượng reach'] 
		tab3.header(headings3)
		lines3 = []
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			ratio = self.get_promotion_ratio(link)
			l3 = []
			l3.append(str(name[:10]))
			l3.append(self.main_sheet.cell(index, type_index).value)
			l3.append(ratio['audience'][0])
			l3.append(ratio['watch_time'][0])
			l3.append(ratio['reach'][0])
			lines3.append(l3)
		for row in range(len(lines3)):
			tab3.add_row(lines3[row])
		s3 = tab3.draw()
		print(s3)		
		'''print("Post link\tThời lượng trung bình\tTrung bình/Full\tTỉ lệ xem hết 95%")
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			n1 = self.get_average_duration(link)
			n2 = self.get_duration_ratio(link)
			n3 = self.get_audience_ratio(link)
			print(str(name[:10]) + "\t%.4f\t%.4f\t%.4f" % (n1, n2, n3))
			print("%.4f" % (n1))
		print()
		print("II. Theo thời gian post")
		print("Post link\tThời gian post\tTỉ lệ like\tTỉ lệ share\tTỉ lệ comment")
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			time = self.main_sheet.cell(index, time_index).value
			ratio = self.get_reaction_ratio(link)
			print(str(name[:10]) + "\t%s\t%.4f\t%.4f\t%.4f" % (time, ratio['like'], ratio['share'], ratio['comment']))
		print()
		print("III. Hiệu quả của promotion")
		print("Post link\tLượng khán gia\tLượt xem\tLượng reach")
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			time = self.main_sheet.cell(index, time_index).value
			ratio = self.get_reaction_ratio(link)
			print(str(name[:10]) + "\t%.4f\t%.4f\t%.4f" % (ratio['like'], ratio['share'], ratio['comment']))'''

	def export_docx(self):
		post_index = self.get_col_index("Post Message", self.main_sheet)
		link_index = self.get_col_index("Permalink", self.main_sheet)
		time_index = self.get_col_index("Posted", self.main_sheet)
		type_index = self.get_col_index("Type", self.main_sheet)

		document = Document()

		document.add_heading('LOK POST PERFORMANCE WEEKLY REPORT', 0)
		p1 = document.add_paragraph('từ ')
		p1.add_run(self.date1 + '/' + self.month).bold = True
		p1.add_run(' đến ')
		p1.add_run(self.date2 + '/' + self.month).bold = True

		document.add_heading('I. Theo thời lượng của post', level = 1)
		table1 = document.add_table(1, cols = 5, style = 'Table Grid')
		hd1 = table1.rows[0].cells
		headings1 = ['Post name', 'Hình thức', 'Thời lượng trung bình', 'Trung bình/Full', 'Tỉ lệ xem hết 95%']
		for i in range(5):
			hd1[i].text = headings1[i]
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			n1 = self.get_average_duration(link)
			n2 = self.get_duration_ratio(link)
			n3 = self.get_audience_ratio(link)
			l1 = []
			l1.append(str(name[:10]))
			l1.append(self.main_sheet.cell(index, type_index).value)
			l1.append(n1)
			l1.append(n2)
			l1.append(n3)
			row_cells = table1.add_row().cells
			for j in range(5):
				row_cells[j].text = str(l1[j])

		document.add_heading('II. Theo thời gian post', level = 1)
		table2 = document.add_table(1, cols = 6, style = 'Table Grid')
		hd2 = table2.rows[0].cells
		headings2 = ['Post name', 'Hình thức', 'Thời gian post', 'Tỉ lệ like', 'Tỉ lệ share', 'Tỉ lệ comment']
		for i in range(6):
			hd2[i].text = headings2[i]	
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			time = (xlrd.xldate_as_tuple(self.main_sheet.cell(index, time_index).value, self.file.datemode))
			ratio = self.get_reaction_ratio(link)
			l2 = []
			l2.append(str(name[:10]))
			l2.append(self.main_sheet.cell(index, type_index).value)
			l2.append(time)
			l2.append(ratio['like'])
			l2.append(ratio['share'])
			l2.append(ratio['comment'])
			row_cells = table2.add_row().cells			
			for j in range(6):
				row_cells[j].text = str(l2[j])

		document.add_heading('III. Hiệu quả của promotion', level = 1)
		table3 = document.add_table(1, cols = 5, style = 'Table Grid')
		hd3 = table3.rows[0].cells
		headings3 = ['Post name', 'Hình thức', 'Lượng khán giả', 'Lượt xem', 'Lượng reach'] 
		for i in range(5):
			hd3[i].text = headings3[i]
		for index in range(2, self.main_sheet.nrows):
			link = self.main_sheet.cell(index, link_index).value
			name = self.main_sheet.cell(index, post_index).value
			ratio = self.get_promotion_ratio(link)
			l3 = []
			l3.append(str(name[:10]))
			l3.append(self.main_sheet.cell(index, type_index).value)
			l3.append(ratio['audience'][0])
			l3.append(ratio['watch_time'][0])
			l3.append(ratio['reach'][0])
			row_cells = table3.add_row().cells
			for j in range(5):
				row_cells[j].text = str(l3[j])

		document.save('/Users/viethongnguyen/Desktop/CSSubmissions/Test/report.docx')







def main():
	test = LACDA_Core()
	test.export_docx()
	#test.print()
	
	'''print(test.get_like("https://www.facebook.com/lokvietnam/videos/1911244905754823/"))
	print(test.get_share("https://www.facebook.com/lokvietnam/videos/1911244905754823/"))
	print(test.get_like("https://www.facebook.com/lokvietnam/videos/1900663656812948/"))
	print(test.get_reach("https://www.facebook.com/lokvietnam/videos/1900663656812948/"))
	print(test.get_reaction_ratio("https://www.facebook.com/lokvietnam/videos/1900663656812948/"))
	print(test.get_audience_ratio("https://www.facebook.com/lokvietnam/videos/1900663656812948/"))
	print(test.get_average_duration("https://www.facebook.com/lokvietnam/videos/1911244905754823/"))
	print(test.get_duration_ratio("https://www.facebook.com/lokvietnam/videos/1911244905754823/"))
	print(test.get_promotion_ratio("https://www.facebook.com/lokvietnam/videos/1911244905754823/"))'''

main()
